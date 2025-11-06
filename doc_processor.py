import io
import docx
import msoffcrypto
import pikepdf
import pytesseract
import camelot
from pdf2image import convert_from_bytes
from pikepdf import PasswordError as PikepdfPasswordError

def _read_docx_text(file_stream, password=None):
    """Extracts text from a DOCX, decrypting if needed, and reads tables."""
    try:
        doc = None
        # 1. Check for encryption
        office_file = msoffcrypto.OfficeFile(file_stream)
        if office_file.is_encrypted():
            if not password:
                raise PermissionError("File is password-protected.")
            
            # 2. Try to decrypt in memory
            decrypted_stream = io.BytesIO()
            office_file.load_key(password=password)
            office_file.decrypt(decrypted_stream)
            doc = docx.Document(decrypted_stream)
        
        else:
            # 3. File is not encrypted, open it directly
            file_stream.seek(0)
            doc = docx.Document(file_stream)
        
        # 4. Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
            
        # 5. Extract text from tables
        if doc.tables:
            text_parts.append("\n\n--- Tables Found ---\n")
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
                text_parts.append("--- End of Table ---\n")
                
        return "\n".join(text_parts)
        
    except msoffcrypto.exceptions.InvalidPassword:
        raise PermissionError("Invalid Password")
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        # This is where the old bug was. We MUST re-raise the error.
        raise e

def _read_pdf_text(file_stream, password=None):
    """
    Complex PDF parser. Tries text, then tables, then OCR.
    """
    text_parts = []
    
    # --- Strategy 1: Try to read text directly with pikepdf ---
    try:
        with pikepdf.open(file_stream, password=password) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text())
        
        full_text = "\n".join(text_parts).strip()
        
        # If we got text, we're done.
        if full_text and len(full_text) > 10:
             print("PDF Strategy: Success (Text Extract)")
             return full_text
             
    except PikepdfPasswordError:
        print("pikepdf: Invalid Password")
        raise PermissionError("Invalid Password") # Re-raise
    except Exception as e:
        # This error (like 'extract_text') means it's not a text-based PDF.
        print(f"PDF Strategy: Text extract failed ({e}). Trying tables.")
        pass # Move to the next strategy
        
    # --- Strategy 2: Try to read tables with Camelot ---
    # We need to reset the stream for each library
    file_stream.seek(0)
    try:
        # Camelot needs the password in a specific way
        tables = camelot.read_pdf(file_stream, password=password, flavor='lattice', pages='all')
        if tables.n > 0:
            print(f"PDF Strategy: Success (Found {tables.n} Tables)")
            text_parts.append("\n\n--- Tables Found ---\n")
            for table in tables:
                text_parts.append(table.df.to_string())
            return "\n".join(text_parts)
        else:
            print("PDF Strategy: Table extract failed (no tables found). Trying OCR.")
    except Exception as e:
        print(f"PDF Strategy: Table extract failed ({e}). Trying OCR.")
        pass # Move to the next strategy

    # --- Strategy 3: Try to read with OCR (scanned image) ---
    file_stream.seek(0)
    try:
        # pdf2image needs poppler-utils (installed in Dockerfile)
        images = convert_from_bytes(file_stream.read())
        print(f"PDF Strategy: Success (OCR on {len(images)} pages)")
        for image in images:
            # pytesseract needs tesseract-ocr (installed in Dockerfile)
            text_parts.append(pytesseract.image_to_string(image))
        
        full_text = "\n".join(text_parts).strip()
        if full_text:
            return full_text
            
    except Exception as e:
        print(f"PDF Strategy: OCR failed ({e}).")
        pass # All strategies failed

    # If all 3 strategies failed, raise an error.
    raise Exception("Could not extract any text, tables, or image data from PDF.")


def extract_document_text(file_data: bytes, filename: str, password: str = None) -> str:
    """
    Main function to extract text from a downloaded file.
    Handles routing to the correct parser (PDF or DOCX).
    """
    file_stream = io.BytesIO(file_data)
    try:
        if filename.lower().endswith('.pdf'):
            return _read_pdf_text(file_stream, password)
        elif filename.lower().endswith('.docx'):
            return _read_docx_text(file_stream, password)
        else:
            return "[Unsupported file type. Can only read .pdf and .docx]"
            
    # --- This block now catches ALL password/decryption errors ---
    except (PermissionError, PikepdfPasswordError):
        raise PermissionError("Invalid Password")
    except Exception as e:
        print(f"Error in extract_document_text: {e}")
        # This will catch the "Could not extract any text" error from PDF
        raise e
    finally:
        file_stream.close()

def check_if_encrypted(file_data: bytes, filename: str) -> bool:
    """
    Quickly checks if a file is encrypted.
    """
    file_stream = io.BytesIO(file_data)
    try:
        if filename.lower().endswith('.pdf'):
            with pikepdf.open(file_stream):
                pass
            return False
        elif filename.lower().endswith('.docx'):
            office_file = msoffcrypto.OfficeFile(file_stream)
            return office_file.is_encrypted()
            
    except PikepdfPasswordError:
        # If it fails with a PasswordError, it IS encrypted
        return True
    except Exception:
        # Corrupt file or other issue
        return False
    finally:
        file_stream.close()